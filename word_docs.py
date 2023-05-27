#python -m pip install python-docx
import docx
wordfile = docx.Document('demo.docx')
print(wordfile.paragraphs[0].text)
#get paragraph text
p = wordfile.paragraphs[1]
print(p.runs)
print(p.runs[3].text)
print(p.runs[3].italic)
#change to underline too
p.runs[3].underline = True
p.runs[3].text = 'Italic and Underlined'
p.style = 'Title'
wordfile.save('demo2.docx')

#starting from blank
new_doc = docx.Document()
new_doc.add_paragraph('The is a paragraph.')
new_doc.add_paragraph('The is another paragraph.')
new_doc.save('demo3.docx')

#add new sentence/run.
pp = new_doc.paragraphs[0]
pp.add_run(' This was added later.')
pp.runs[1].bold = True
new_doc.save('demo3.docx')